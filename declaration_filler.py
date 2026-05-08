import json
import logging
import re
from datetime import datetime
from pathlib import Path

import anthropic
import fitz  # PyMuPDF
from docx import Document
from docx.shared import Pt

from config import ANTHROPIC_API_KEY, DATOS_EMPRESA, MODEL_FAST, TEMPLATES_DIR

logger = logging.getLogger(__name__)
client = anthropic.Anthropic(api_key=ANTHROPIC_API_KEY)

INDEX_PATH = TEMPLATES_DIR / "declaraciones" / "index.json"


# ── Text extraction helpers ───────────────────────────────────────────────────

def _extract_page_text(pdf_path: str, page_num: int) -> str:
    try:
        doc = fitz.open(pdf_path)
        idx = max(0, page_num - 1)
        pages = [doc[i].get_text() for i in range(max(0, idx - 1), min(len(doc), idx + 3))]
        doc.close()
        return "\n".join(pages)
    except Exception as e:
        logger.warning(f"Error extrayendo página {page_num} de {pdf_path}: {e}")
        return ""


def _extract_docx_text(docx_path: str) -> str:
    doc = Document(docx_path)
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text for cell in row.cells))
    return "\n".join(line for line in parts if line.strip())


def _contract_context(analysis: dict) -> str:
    ident = analysis.get("identificacion", {})
    return (
        f"Expediente: {ident.get('expediente', 'N/A')}\n"
        f"Objeto: {ident.get('objeto', 'N/A')}\n"
        f"Órgano de contratación: {ident.get('organo_contratacion', 'N/A')}\n"
        f"Tipo de contrato: {ident.get('tipo_contrato', 'N/A')}\n"
        f"Procedimiento: {ident.get('procedimiento', 'N/A')}\n"
        f"Presupuesto base (sin IVA): {ident.get('presupuesto_base', 'N/A')}\n"
        f"Valor estimado: {ident.get('valor_estimado', 'N/A')}\n"
        f"Duración: {ident.get('duracion', 'N/A')}\n"
    )


def _save_docx(filled_text: str, output_path: str, title: str) -> None:
    doc = Document()
    doc.add_heading(title, level=1)
    for line in filled_text.split("\n"):
        p = doc.add_paragraph(line)
        for run in p.runs:
            run.font.size = Pt(11)
    doc.save(output_path)


def _fill_docx_preserving_format(template_path: str, filled_text: str, out_path: str) -> None:
    doc = Document(template_path)
    filled_lines = [l for l in filled_text.split("\n")]
    line_iter = iter(filled_lines)
    for para in doc.paragraphs:
        if not para.text.strip():
            continue
        new_line = next(line_iter, None)
        if new_line is None:
            break
        if para.runs:
            style = para.runs[0]
            for run in para.runs:
                run.text = ""
            style.text = new_line
        else:
            para.text = new_line
    doc.save(out_path)


def _extract_json(text: str) -> dict | list:
    m = re.search(r"```(?:json)?\s*(\{.*?\}|\[.*?\])\s*```", text, re.DOTALL)
    if m:
        return json.loads(m.group(1))
    text = text.strip()
    if text.startswith("{") or text.startswith("["):
        return json.loads(text)
    for start_char, end_char in (("{", "}"), ("[", "]")):
        s, e = text.find(start_char), text.rfind(end_char)
        if s != -1 and e != -1:
            return json.loads(text[s:e + 1])
    raise ValueError("No se encontró JSON en la respuesta")


# ── INDEXING (one-time, run with --indexar-plantillas) ────────────────────────

def _index_single(tpl_path: Path) -> dict | None:
    """Ask Claude to summarise one template in a compact index entry."""
    try:
        text = _extract_docx_text(str(tpl_path))
    except Exception as e:
        logger.error(f"Indexado: no se pudo leer '{tpl_path.name}': {e}")
        return None

    prompt = (
        f"Analiza esta plantilla de declaración para licitaciones públicas españolas "
        f"y devuelve ÚNICAMENTE este JSON (sin texto extra):\n"
        f'{{"titulo": "nombre descriptivo corto", '
        f'"descripcion": "para qué sirve esta declaración en 1-2 frases", '
        f'"aplica_tipos": ["suministro", "servicios", "obras"], '
        f'"aplica_siempre": true, '
        f'"no_aplica_si": "condición en que NO aplica o vacío si siempre aplica", '
        f'"campos_clave": ["campo1", "campo2"]}}\n\n'
        f"PLANTILLA:\n{text[:4000]}"
    )
    with client.messages.stream(
        model=MODEL_FAST,
        max_tokens=512,
        messages=[{"role": "user", "content": prompt}],
    ) as stream:
        msg = stream.get_final_message()

    raw = "".join(b.text for b in msg.content if hasattr(b, "text"))
    try:
        entry = _extract_json(raw)
        entry["filename"] = tpl_path.name
        return entry
    except Exception as e:
        logger.error(f"Indexado: JSON inválido para '{tpl_path.name}': {e}")
        return None


def index_templates(force: bool = False) -> None:
    """
    One-time indexing of all DOCX templates in templates/declaraciones/.
    Run once after adding new templates:  python main.py --indexar-plantillas
    """
    templates_dir = TEMPLATES_DIR / "declaraciones"
    all_files = sorted(templates_dir.glob("*.docx"))
    if not all_files:
        print("No se encontraron archivos .docx en templates/declaraciones/")
        return

    # Load existing index
    existing: dict = {}
    if INDEX_PATH.exists() and not force:
        try:
            data = json.loads(INDEX_PATH.read_text(encoding="utf-8"))
            existing = {e["filename"]: e for e in data.get("plantillas", [])}
        except Exception:
            pass

    to_index = [f for f in all_files if force or f.name not in existing]
    if not to_index:
        print(f"Índice ya actualizado ({len(existing)} plantillas). "
              f"Usa --forzar para reindexar todo.")
        return

    print(f"Indexando {len(to_index)} plantilla(s)…")
    for tpl in to_index:
        print(f"  · {tpl.name}… ", end="", flush=True)
        entry = _index_single(tpl)
        if entry:
            existing[tpl.name] = entry
            print("OK")
        else:
            print("ERROR (se omite)")

    # Remove entries for files that no longer exist
    existing = {k: v for k, v in existing.items()
                if (templates_dir / k).exists()}

    payload = {
        "version": 1,
        "indexed_at": datetime.now().isoformat(timespec="seconds"),
        "plantillas": list(existing.values()),
    }
    INDEX_PATH.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
    print(f"\n✔ Índice guardado: {INDEX_PATH}  ({len(existing)} plantillas)")


# ── PASO 5a: declarations from the pliego itself ─────────────────────────────

def _fill_with_claude(template_text: str, campos: list[str], analysis: dict) -> str:
    prompt = (
        f"Eres un especialista en licitaciones públicas españolas.\n\n"
        f"DATOS DE LA EMPRESA:\n{json.dumps(DATOS_EMPRESA, ensure_ascii=False, indent=2)}\n\n"
        f"DATOS DEL CONTRATO:\n{_contract_context(analysis)}\n"
        f"CAMPOS A COMPLETAR:\n{json.dumps(campos, ensure_ascii=False)}\n\n"
        f"TEXTO ORIGINAL DE LA DECLARACIÓN:\n{template_text}\n\n"
        f"Devuelve el texto completo con todos los campos en blanco rellenados. "
        f"Mantén el formato original. Sin explicaciones adicionales."
    )
    with client.messages.stream(
        model=MODEL_FAST,
        max_tokens=4096,
        messages=[{"role": "user", "content": prompt}],
    ) as stream:
        msg = stream.get_final_message()
    return "".join(b.text for b in msg.content if hasattr(b, "text"))


def fill_declarations(
    declaraciones: list[dict],
    file_map: dict,
    analysis: dict,
    output_dir: str,
) -> list[str]:
    output_path = Path(output_dir)
    results = []

    for decl in declaraciones:
        nombre = decl.get("nombre", "declaracion")
        doc_origen = decl.get("documento_origen", "")
        pagina = int(decl.get("pagina", 1) or 1)
        campos = decl.get("campos_a_rellenar", [])

        logger.info(f"PASO 5a: Rellenando '{nombre}' desde '{doc_origen}'…")

        template_text = ""
        pdf_path = file_map.get(doc_origen)
        if pdf_path and Path(pdf_path).suffix.lower() == ".pdf":
            template_text = _extract_page_text(pdf_path, pagina)
        if not template_text:
            logger.warning(f"PASO 5a: Sin texto para '{nombre}', usando descripción")
            template_text = decl.get("descripcion", nombre)

        try:
            filled = _fill_with_claude(template_text, campos, analysis)
        except Exception as e:
            logger.error(f"PASO 5a: Error rellenando '{nombre}': {e}")
            continue

        safe_name = re.sub(r'[^\w\-]', '_', nombre)[:60]
        docx_path = str(output_path / f"declaracion_{safe_name}.docx")
        try:
            _save_docx(filled, docx_path, nombre)
            logger.info(f"PASO 5a: Guardado {docx_path}")
            results.append(docx_path)
        except Exception as e:
            logger.error(f"PASO 5a: Error guardando DOCX '{nombre}': {e}")

    return results


# ── PASO 5b: generic user templates (index-based, token-efficient) ────────────

def _select_applicable(
    index_entries: list[dict],
    analysis: dict,
    declaraciones_del_pliego: list[dict],
) -> list[dict]:
    """
    Lightweight call: send only the index (not template content) to Claude
    and get back which filenames apply to this contract.
    Templates are only selected when the pliego has no explicit annex
    covering the same declaration type.
    """
    compact = [
        {
            "filename": e["filename"],
            "titulo": e.get("titulo", ""),
            "descripcion": e.get("descripcion", ""),
            "aplica_tipos": e.get("aplica_tipos", []),
            "aplica_siempre": e.get("aplica_siempre", True),
            "no_aplica_si": e.get("no_aplica_si", ""),
        }
        for e in index_entries
    ]

    # Compact list of what the pliego already provides
    ya_cubiertos = [
        {"nombre": d.get("nombre", ""), "descripcion": d.get("descripcion", "")}
        for d in declaraciones_del_pliego
    ]

    ya_cubiertos_txt = (
        f"DECLARACIONES YA INCLUIDAS EN EL PROPIO PLIEGO (NO usar plantilla genérica para estas):\n"
        f"{json.dumps(ya_cubiertos, ensure_ascii=False, indent=2)}\n\n"
        if ya_cubiertos else
        "El pliego no incluye ningún anexo/declaración propio.\n\n"
    )

    prompt = (
        f"DATOS DEL CONTRATO:\n{_contract_context(analysis)}\n\n"
        f"{ya_cubiertos_txt}"
        f"PLANTILLAS GENÉRICAS DISPONIBLES (de concursos anteriores):\n"
        f"{json.dumps(compact, ensure_ascii=False, indent=2)}\n\n"
        f"Selecciona SOLO las plantillas que:\n"
        f"  1. Aplican al tipo de contrato y procedimiento de este concurso, Y\n"
        f"  2. NO están ya cubiertas por un anexo explícito del propio pliego.\n"
        f'Devuelve ÚNICAMENTE este JSON:\n'
        f'{{"aplican": ["filename1.docx", "filename2.docx"]}}'
    )
    with client.messages.stream(
        model=MODEL_FAST,
        max_tokens=256,
        messages=[{"role": "user", "content": prompt}],
    ) as stream:
        msg = stream.get_final_message()

    raw = "".join(b.text for b in msg.content if hasattr(b, "text"))
    try:
        result = _extract_json(raw)
        selected_names = set(result.get("aplican", []))
        return [e for e in index_entries if e["filename"] in selected_names]
    except Exception as e:
        logger.error(f"PASO 5b: Error parseando selección: {e}")
        return []


def _fill_single_template(entry: dict, analysis: dict, output_dir: str) -> str | None:
    """Fill one template (send content only for this one file)."""
    tpl_path = TEMPLATES_DIR / "declaraciones" / entry["filename"]
    if not tpl_path.exists():
        logger.warning(f"PASO 5b: Plantilla no encontrada: {tpl_path}")
        return None

    try:
        text = _extract_docx_text(str(tpl_path))
    except Exception as e:
        logger.error(f"PASO 5b: No se pudo leer '{entry['filename']}': {e}")
        return None

    campos = entry.get("campos_clave", [])
    prompt = (
        f"Eres un especialista en licitaciones públicas españolas.\n\n"
        f"DATOS DE LA EMPRESA:\n{json.dumps(DATOS_EMPRESA, ensure_ascii=False, indent=2)}\n\n"
        f"DATOS DEL CONTRATO:\n{_contract_context(analysis)}\n"
        f"CAMPOS ESPERADOS: {json.dumps(campos, ensure_ascii=False)}\n\n"
        f"TEXTO DE LA DECLARACIÓN:\n{text}\n\n"
        f"Devuelve el texto completo rellenado con los datos correctos. "
        f"Mantén el formato. Sin explicaciones."
    )
    with client.messages.stream(
        model=MODEL_FAST,
        max_tokens=4096,
        messages=[{"role": "user", "content": prompt}],
    ) as stream:
        msg = stream.get_final_message()

    filled = "".join(b.text for b in msg.content if hasattr(b, "text"))

    safe_name = re.sub(r'[^\w\-]', '_', Path(entry["filename"]).stem)[:60]
    out_path = str(Path(output_dir) / f"plantilla_{safe_name}_rellenada.docx")
    try:
        _fill_docx_preserving_format(str(tpl_path), filled, out_path)
    except Exception:
        _save_docx(filled, out_path, entry.get("titulo", entry["filename"]))

    return out_path


def fill_from_templates(
    analysis: dict,
    output_dir: str,
    declaraciones_del_pliego: list[dict] | None = None,
) -> list[str]:
    """
    Token-efficient flow:
      1. Read the lightweight index (no file content) — ~150 tokens
      2. Ask Claude which templates apply AND are not already covered
         by an explicit annex in the pliego — tiny call
      3. Send content only for the selected templates — typically 2-3 files
    """
    if not INDEX_PATH.exists():
        logger.warning(
            "PASO 5b: No hay índice de plantillas. "
            "Añade tus DOCX a templates/declaraciones/ y ejecuta: "
            "python main.py --indexar-plantillas"
        )
        return []

    try:
        data = json.loads(INDEX_PATH.read_text(encoding="utf-8"))
    except Exception as e:
        logger.error(f"PASO 5b: Error leyendo índice: {e}")
        return []

    entries = data.get("plantillas", [])
    if not entries:
        logger.info("PASO 5b: Índice vacío")
        return []

    pliego_decls = declaraciones_del_pliego or []
    if pliego_decls:
        logger.info(
            f"PASO 5b: El pliego ya cubre {len(pliego_decls)} declaración(es) — "
            f"las plantillas genéricas solo cubrirán el resto"
        )
    logger.info(f"PASO 5b: Seleccionando entre {len(entries)} plantillas indexadas…")
    selected = _select_applicable(entries, analysis, pliego_decls)

    if not selected:
        logger.info("PASO 5b: Ninguna plantilla genérica aplica a este contrato")
        return []

    logger.info(f"PASO 5b: {len(selected)} plantilla(s) seleccionada(s): "
                f"{[e['filename'] for e in selected]}")

    results = []
    for entry in selected:
        logger.info(f"PASO 5b: Rellenando '{entry['filename']}'…")
        out = _fill_single_template(entry, analysis, output_dir)
        if out:
            results.append(out)
            logger.info(f"PASO 5b: Guardado {out}")
        else:
            logger.error(f"PASO 5b: Falló el relleno de '{entry['filename']}'")

    return results
